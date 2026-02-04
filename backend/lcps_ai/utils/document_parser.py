"""
Document parsing utilities for PDF and Word documents
"""
import os
from pathlib import Path
from typing import Optional
import PyPDF2
import pdfplumber
from docx import Document
from utils.logger import logger


class DocumentParser:
    """Parse PDF and Word documents"""
    
    @staticmethod
    def parse_pdf(file_path: str, method: str = 'pypdf2') -> Optional[str]:
        """
        Parse PDF file and extract text
        
        Args:
            file_path: Path to PDF file
            method: 'pypdf2' or 'pdfplumber' (pdfplumber is better for complex PDFs)
        
        Returns:
            Extracted text or None if parsing fails
        """
        if not os.path.exists(file_path):
            logger.error(f"PDF file not found: {file_path}")
            return None
        
        try:
            if method == 'pdfplumber':
                return DocumentParser._parse_with_pdfplumber(file_path)
            else:
                return DocumentParser._parse_with_pypdf2(file_path)
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            return None
    
    @staticmethod
    def _parse_with_pypdf2(file_path: str) -> str:
        """Parse PDF using PyPDF2"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    @staticmethod
    def _parse_with_pdfplumber(file_path: str) -> str:
        """Parse PDF using pdfplumber (better for complex PDFs)"""
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    
    @staticmethod
    def parse_docx(file_path: str) -> Optional[str]:
        """
        Parse Word document and extract text
        
        Args:
            file_path: Path to Word document
        
        Returns:
            Extracted text or None if parsing fails
        """
        if not os.path.exists(file_path):
            logger.error(f"Word document not found: {file_path}")
            return None
        
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing Word document: {e}")
            return None
    
    @staticmethod
    def parse_document(file_path: str) -> Optional[str]:
        """
        Auto-detect document type and parse
        
        Args:
            file_path: Path to document
        
        Returns:
            Extracted text or None if parsing fails
        """
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            # Try pdfplumber first, fallback to PyPDF2
            text = DocumentParser.parse_pdf(file_path, method='pdfplumber')
            if not text or len(text) < 50:  # If extraction failed or too short
                text = DocumentParser.parse_pdf(file_path, method='pypdf2')
            return text
        elif file_extension in ['.docx', '.doc']:
            return DocumentParser.parse_docx(file_path)
        else:
            logger.error(f"Unsupported file format: {file_extension}")
            return None
    
    @staticmethod
    def get_document_info(file_path: str) -> dict:
        """Get basic document information"""
        if not os.path.exists(file_path):
            return {}
        
        file_stat = os.stat(file_path)
        return {
            'filename': os.path.basename(file_path),
            'size': file_stat.st_size,
            'extension': Path(file_path).suffix.lower(),
            'modified': file_stat.st_mtime
        }
